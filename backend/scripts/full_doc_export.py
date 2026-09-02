import asyncio
import json
from motor.motor_asyncio import AsyncIOMotorClient
from backend.database import db
from bson import ObjectId
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

async def main():
    db.connect()
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # =====================================================================
        # TITLE
        # =====================================================================
        title = doc.add_heading('Questionnaire Platform — Complete Question & Attribute Reference', level=0)
        doc.add_paragraph(
            'This document contains:\n'
            '  1. All Attribute Banks (Main vs Sub attributes across the platform)\n'
            '  2. All Module Questions (Purchase Funnel, Brand Usage, Brand Pricing Behavior)\n'
            '  3. Taste Test Evaluation Structure (scale scores, open-ended questions)\n'
            '  4. Demographics Fields\n'
        )

        # =====================================================================
        # PART 1 — ATTRIBUTE BANKS
        # =====================================================================
        doc.add_heading('Part 1: Attribute Banks (Main vs Sub)', level=1)
        doc.add_paragraph(
            'The Attribute Bank is the master library of sensory/diagnostic attributes. '
            'Each bank has a category (e.g. "appearance", "texture") and contains '
            'CORE (Main) Attributes and SUB Attributes. Main attributes use a '
            'broader "Linear Scale 1-10", while Sub attributes use a finer '
            '"Scale 1-5" for deeper diagnostics.'
        )

        banks = await db.get_collection("attribute_banks").find({}).to_list(length=None)
        for bank in banks:
            category = bank.get("category", "unknown")
            display = bank.get("display_name", category)
            doc.add_heading(f'{display} ({category})', level=2)

            core = bank.get("core_attributes", [])
            sub = bank.get("sub_attributes", [])

            if core:
                doc.add_heading('Main (Core) Attributes', level=3)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Attribute ID'
                hdr[1].text = 'Label'
                hdr[2].text = 'Scale Type'
                for a in core:
                    row = table.add_row().cells
                    row[0].text = a.get('attribute_id', '')
                    row[1].text = a.get('label', '')
                    row[2].text = a.get('scale_type', '')

            if sub:
                doc.add_heading('Sub Attributes', level=3)
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Attribute ID'
                hdr[1].text = 'Label'
                hdr[2].text = 'Scale Type'
                for a in sub:
                    row = table.add_row().cells
                    row[0].text = a.get('attribute_id', '')
                    row[1].text = a.get('label', '')
                    row[2].text = a.get('scale_type', '')
            
            if not sub:
                doc.add_paragraph('(No sub attributes defined for this category)')

        # =====================================================================
        # PART 2 — QUESTION MODULES
        # =====================================================================
        doc.add_heading('Part 2: Question Modules', level=1)
        
        modules = await db.get_collection("question_modules").find({}).to_list(length=None)
        for mod in modules:
            mod_name = mod.get("name", "Unknown")
            mod_desc = mod.get("description", "")
            doc.add_heading(mod_name, level=2)
            doc.add_paragraph(mod_desc)

            for sec in mod.get("sections", []):
                sec_title = sec.get("title_en", sec.get("section_id", ""))
                doc.add_heading(f'Section: {sec_title}', level=3)

                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Q ID'
                hdr[1].text = 'Label'
                hdr[2].text = 'Type'
                hdr[3].text = 'English Text'
                hdr[4].text = 'Arabic Text'

                for q in sec.get("questions", []):
                    row = table.add_row().cells
                    row[0].text = q.get("question_id", "")
                    row[1].text = q.get("label", "")
                    row[2].text = q.get("type", "")
                    row[3].text = q.get("en_text", "")
                    row[4].text = q.get("ar_text", "")

        # =====================================================================
        # PART 3 — TASTE TEST EVALUATION STRUCTURE
        # =====================================================================
        doc.add_heading('Part 3: Taste Test Evaluation Structure', level=1)
        doc.add_paragraph(
            'In Taste Test surveys, each respondent evaluates each brand (rotation) '
            'on the following scale attributes and open-ended questions.'
        )

        doc.add_heading('Scale Scores (26 Attributes for Hero Protein Bar)', level=2)
        scores_list = [
            ("Overall Likeness", "1-9", "Hedonic (overall liking)"),
            ("Outershape", "1-5", "Diagnostic"),
            ("Outercolor", "1-5", "Diagnostic"),
            ("Smell", "1-5", "Diagnostic"),
            ("Natural Smell", "1-5", "Diagnostic"),
            ("Innercolor", "1-5", "Diagnostic"),
            ("Natural Inside", "1-5", "Diagnostic"),
            ("Texture", "1-5", "Diagnostic"),
            ("Consistency", "1-5", "Diagnostic"),
            ("Firmness", "1-5", "Diagnostic"),
            ("Freshness", "1-5", "Diagnostic"),
            ("Size", "1-5", "Diagnostic"),
            ("General Shape", "1-5", "Diagnostic"),
            ("Ease of Bite", "1-5", "Diagnostic"),
            ("Chewing Experience", "1-5", "Diagnostic"),
            ("Ease of Swallow", "1-5", "Diagnostic"),
            ("Eating Experience", "1-5", "Diagnostic"),
            ("Natural Taste", "1-5", "Diagnostic"),
            ("Taste Intensity", "1-5", "Diagnostic"),
            ("Sweetness", "1-5", "Diagnostic"),
            ("Taste Consistency", "1-5", "Diagnostic"),
            ("After Taste", "1-5", "Diagnostic"),
            ("Bitterness", "1-5", "Diagnostic"),
            ("After Taste Duration", "1-5", "Diagnostic"),
            ("Taste Quality", "1-5", "Diagnostic"),
            ("Essence", "1-9", "Hedonic"),
        ]
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '#'
        hdr[1].text = 'Attribute Name'
        hdr[2].text = 'Scale'
        hdr[3].text = 'Type'
        for i, (name, scale, atype) in enumerate(scores_list, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = name
            row[2].text = scale
            row[3].text = atype

        doc.add_heading('Open-Ended Questions (6 per brand rotation)', level=2)
        oe_table = doc.add_table(rows=1, cols=2)
        oe_table.style = 'Table Grid'
        hdr = oe_table.rows[0].cells
        hdr[0].text = 'Field Key'
        hdr[1].text = 'Question'
        oe_items = [
            ("favorite_shape", "What did you like most about the shape/appearance?"),
            ("worst_shape", "What did you dislike most about the shape/appearance?"),
            ("improvement_shape", "What improvements would you suggest for the shape/appearance?"),
            ("favorite_taste", "What did you like most about the taste?"),
            ("worst_taste", "What did you dislike most about the taste?"),
            ("improvement_taste", "What improvements would you suggest for the taste?"),
        ]
        for key, desc in oe_items:
            row = oe_table.add_row().cells
            row[0].text = key
            row[1].text = desc

        doc.add_heading('Additional Taste Test Fields', level=2)
        doc.add_paragraph('flavor — What flavor was tested')
        doc.add_paragraph('preferred_brand — Which brand the respondent preferred overall')
        doc.add_paragraph('preferred_brand_other — Why they preferred it (open text)')
        doc.add_paragraph('purchase_price — How much would respondent pay for this brand (numeric)')

        # =====================================================================
        # PART 4 — DEMOGRAPHICS
        # =====================================================================
        doc.add_heading('Part 4: Demographics Module', level=1)
        demo_table = doc.add_table(rows=1, cols=3)
        demo_table.style = 'Table Grid'
        hdr = demo_table.rows[0].cells
        hdr[0].text = 'Field'
        hdr[1].text = 'Type'
        hdr[2].text = 'Example Values'
        demo_items = [
            ("name", "Text", "ماجد جورج"),
            ("phone", "Text", "1221702417"),
            ("gender", "Choice", "male / female"),
            ("age_group", "Choice", "18-24, 25-34, 35-44, 45-54, 55+"),
        ]
        for field, ftype, example in demo_items:
            row = demo_table.add_row().cells
            row[0].text = field
            row[1].text = ftype
            row[2].text = example

        # =====================================================================
        # PART 5 — THIS SURVEY's ATTRIBUTE MAPPING
        # =====================================================================
        doc.add_heading('Part 5: Hero Protein Bar Taste Test — Attribute Mapping', level=1)
        doc.add_paragraph(
            'For this specific survey, every attribute was configured as a '
            'Main attribute with itself as the only Sub attribute. This means '
            'there are NO distinct sub-attributes — each main IS its own sub.'
        )
        survey = await db.get_collection("surveys").find_one({"_id": ObjectId("6a3b8939b3fa5ef1308239ed")})
        config = survey.get("taste_test_config", {})
        attr_seq = config.get("attribute_sequence", [])
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Main Attribute'
        hdr[1].text = 'Sub Attributes'
        for attr in attr_seq:
            row = table.add_row().cells
            row[0].text = attr.get("main_attribute", "")
            row[1].text = ", ".join(attr.get("sub_attributes", []))

        doc.save('/app/backend/Platform_Questions_And_Attributes_Reference.docx')
        print("SUCCESS! Document saved.")
    finally:
        db.close()

if __name__ == '__main__':
    asyncio.run(main())
