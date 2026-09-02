import asyncio
import docx
from motor.motor_asyncio import AsyncIOMotorClient
from backend.database import db
from bson import ObjectId

async def main():
    db.connect()
    try:
        survey = await db.get_collection("surveys").find_one({"_id": ObjectId("6a3b8939b3fa5ef1308239ed")})
        doc = docx.Document()
        doc.add_heading('Survey Questions: Hero Protein Bar Taste Test', 0)
        
        if survey:
            snapshot = survey.get("template_snapshot_l2", {})
            sections = snapshot.get("sections", [])
            for section in sections:
                doc.add_heading(section.get('title', 'Unknown Section'), level=1)
                for q in section.get("questions", []):
                    # Write Question text
                    text = q.get('text', 'No text')
                    doc.add_paragraph(f"Q: {text}", style='List Number')
                    # Write type
                    qtype = q.get('type', '')
                    p = doc.add_paragraph(f"Type: {qtype}")
                    # Write options if any
                    options = q.get('options', [])
                    if options:
                        for opt in options:
                            doc.add_paragraph(f"- {opt.get('text', '')}", style='List Bullet 2')
        else:
            doc.add_paragraph("Survey not found.")
            
        doc.save('/app/backend/Survey_Questions_Export.docx')
        print("Document saved to /app/backend/Survey_Questions_Export.docx")
    finally:
        db.close()

if __name__ == '__main__':
    asyncio.run(main())
